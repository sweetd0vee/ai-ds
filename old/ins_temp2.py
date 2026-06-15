import os
import streamlit as st
import pandas as pd
from io import BytesIO, StringIO
import json
import logging
from dateutil.parser import parse
import numpy as np
import re
import ast
import textwrap

# Настройка логгирования для отладки
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from langchain_community.llms import Ollama
from langchain_core.prompts import PromptTemplate
from langchain_classic.chains import LLMChain
from langchain_experimental.utilities import PythonREPL

# --- ПРОМТЫ ---
struct_analyze = """
                Проанализируй следующую информацию о файле:
                {file_info}
                Определи названия столбцов и их типы данных (числовой, категориальный, дата, булевый и т.д.).
                Особое внимание уделите столбцам, которые могут содержать дату/время.

                Предоставь анализ в формате, строго следуя шаблону ниже. Не добавляй пояснения или другой текст до или после:
                ---COLUMNS_START---
                Столбец: имя_столбца_1
                Тип: предполагаемый_тип_данных
                Описание: краткое описание назначения
                
                Столбец: имя_столбца_2
                Тип: предполагаемый_тип_данных
                Описание: краткое описание назначения
                ---COLUMNS_END---
                ---DATETIME_CANDIDATES_START---
                имя_столбца_1, имя_столбца_3
                ---DATETIME_CANDIDATES_END---

                Пример (не копировать):
                ---COLUMNS_START---
                Столбец: PassengerId
                Тип: numerical (integer)
                Описание: Уникальный идентификатор пассажира

                Столбец: Name
                Тип: textual (object)
                Описание: Имя пассажира
                ---COLUMNS_END---
                ---DATETIME_CANDIDATES_START---
                
                ---DATETIME_CANDIDATES_END---
                """ #Предварительный анализ входных данных

m_plan =  """
                На основе следующего описания данных:
                {data_structure}
                Предложи список статистических метрик, которые стоит рассчитать для анализа данных. Учти Best Practices по анализу данных:

                - Для числовых столбцов:
                    * Меры центральной тенденции: mean (среднее арифметическое), median (медиана), mode (наиболее часто встречающееся значение)
                    * Меры изменчивости: std (стандартное отклонение), var (дисперсия), mad (среднее абсолютное отклонение от среднего)
                    * Меры формы распределения: skew (коэффициент асимметрии), kurtosis (коэффициент эксцесса)
                    * Квантили: min (минимум), max (максимум), quantile_25 (1-й квартиль, 25%), quantile_75 (3-й квартиль, 75%), quantile_90 (90-й перцентиль), quantile_95 (95-й перцентиль)
                    * Другие: count (количество непустых значений), iqr (интерквартильный размах)

                - Для категориальных столбцов:
                    * count (общее количество непустых значений)
                    * nunique (количество уникальных категорий)
                    * mode (наиболее частотная категория)
                    * mode_count (частота (количество) самой частотной категории)
                    * mode_rel_freq (относительная частота (доля) самой частотной категории)

                - Для дат (datetime):
                    * count (количество непустых значений)
                    * min_date (минимальная/ранняя дата)
                    * max_date (максимальная/поздняя дата)
                    * date_range_days (общий диапазон дат в днях, разница max_date - min_date в днях)
                    * unique_dates (количество уникальных дат/дней)
                    * dates_per_month (среднее количество записей на месяц, если данных достаточно)

                Верни список метрик в формате, строго следуя шаблону ниже. Не добавляй пояснения или другой текст до или после:
                ---METRICS_START---
                Столбец: столбец_1
                Метрики: метрика1, метрика2, метрика3

                Столбец: столбец_2
                Метрики: метрика4, метрика5
                ---METRICS_END---

                Пример (не копировать):
                ---METRICS_START---
                Столбец: Age
                Метрики: count, mean, median, mode, std, var, min, max, quantile_25, quantile_75, skew, kurtosis, mad, iqr

                Столбец: Sex
                Метрики: count, nunique, mode, mode_count, mode_rel_freq

                Столбец: SignupDate
                Метрики: count, min_date, max_date, date_range_days, unique_dates
                ---METRICS_END---
                """ #создание плана метрик

# data_analyze = """
# Ниже прведены рассчитанные метрики для каждого столбца:
# {metrics_results_raw}

# Твоя задача - провести глубокий и содержательный анализ этих метрик и предоставить интерпретацию на русском языке.

# 📊 ТРЕБОВАНИЯ К АНАЛИЗУ:
# - Для КАЖДОГО столбца дай отдельный анализ
# - Опиши все ключевые статистические показатели (среднее, медиана, мода, std, мин/макс)
# - Выяви тенденции, закономерности и паттерны
# - Отметь аномалии, выбросы и необычные значения
# - Сравни распределения между столбцами где возможно
# - Сделай выводы о нормальности распределения
# - Укажи потенциальные проблемы с данными (пропуски, дубликаты, ошибки)
# - Дай рекомендации по дальнейшей работе с каждым столбцом
# - Сформулируй гипотезы и предположения на основе данных

# 📌 СТРУКТУРА АНАЛИЗА:
# 1. Общая характеристика датасета
# 2. Подробный анализ каждого столбца по отдельности:
#    - Основные статистики и их интерпретация
#    - Форма распределения и отклонения
#    - Аномалии и выбросы
#    - Практическое значение показателей
# 3. Возможные межстолбцовые связи и зависимости
# 4. Выводы и рекомендации

# Используй профессиональную терминологию анализа данных. Будь конкретен в цифрах и примерах.
# """
# data_analyze = """
# Ниже приведены рассчитанные метрики для каждого столбца:
# {metrics_results_raw}

# Твоя задача - провести ГЛУБОКИЙ, ИСЧЕРПЫВАЮЩИЙ и СОДЕРЖАТЕЛЬНЫЙ анализ этих метрик на русском языке.

# Для каждого столбца рассмотри следующие вопросы
# - Основные статистические характеристики (среднее, медиана, стандартное отклонение), распределение данных (нормальное, скошенное, наличие выбросов)
# - Тенденции и закономерности, аномалии и необычные значения
# - Потенциальные проблемы с данными (пропуски, дубликаты, некорректные значения)
# - Интерпретация значений в контексте предметной области, гипотезы и предположения о причинах наблюдаемых паттернов

# ⚠️ ВАЖНО:
# - Не просто описывай числа - объясняй их смысл
# - Делай акцент на практических выводах
# - Указывай на потенциальные проблемы и возможности
# - Формулируй гипотезы для дальнейшего исследования

# Отчет должен быть ПОЛЕЗНЫМ для аналитика, который будет строить графики и делать выводы на основе этого анализа.
# """
data_analyze =  """
                Ниже приведены рассчитанные метрики для каждого столбца:
                {metrics_results_raw}

                Твоя задача - провести глубокий и содержательный анализ этих метрик и предоставить интерпретацию на русском языке.
Обычно объем отчета - 40-50 предложений.
                """ #Аналитический блок на метриках

final_rep = """
На основе следующего анализа метрик:
{analysis_summary}

Твоя задача — написать **комплексный, содержательный и профессиональный итоговый отчёт** для заинтересованной стороны (например, руководителя отдела аналитики или бизнес-подразделения). Отчёт должен не просто пересказывать цифры, а **рассказывать историю данных**, выделять ключевые инсайты и давать ценные рекомендации.

**Стиль и тон:**
- Пиши **четко, профессионально, но не сухо**. Избегай излишнего канцелярского языка.
- Используй **аналитический и повествовательный стиль**. Покажи, как данные ведут к определенным выводам.
- Будь **конкретен и ориентирован на бизнес-контекст**. Связывай метрики с реальными последствиями (например, "высокое стандартное отклонение в доходности указывает на высокую волатильность и, как следствие, на потенциальный риск для портфеля").
- Фокусируйся на **"почему" и "что это значит"**, а не только на "что".

**Структура и содержание:**
Строго следуй приведенной ниже структуре. Каждая секция должна быть логично связана с предыдущей.

---
==СВОДКА==
Напиши краткое, но мощное вступление (3-4 предложения). Сформулируй **самые важные открытия** из анализа. Ответь на вопрос: "Какие **ключевые проблемы, возможности или тренды** выявили данные?" Этот раздел должен дать читателю мгновенное понимание сути отчёта.

==ИНФОРМАЦИЯ О ДАННЫХ==
Кратко опиши источник данных:
- Общее количество записей и признаков.
- Основные типы данных (числовые, категориальные, дата/время).
- Упомяни **ограничения по качеству данных**, если они были выявлены (например, пропуски, аномалии), и как они могли повлиять на анализ.

==ГЛУБОКИЙ АНАЛИЗ КЛЮЧЕВЫХ ПЕРЕМЕННЫХ==
Представь это как **сердце отчёта**. Для каждой из **3-5 наиболее значимых переменных** (выбери те, которые, по твоему мнению, несут наибольшую аналитическую ценность) предоставь отдельный подраздел. Для менее значимых переменных можно дать краткое резюме в конце этого раздела.

[Имя столбца] ([тип: числовой/категориальный/дата])
- **Ключевые метрики:** Приведи самые важные цифры (среднее, медиана, мода, стандартное отклонение, уникальные значения и т.д.).
- **Интерпретация и контекст:** Объясни, **что означают эти метрики** в бизнес-контексте. Например, "Средняя доходность в 15% при медиане 8% указывает на наличие нескольких клиентов с очень высокой доходностью, которые "перетягивают" среднее вверх, что может скрывать более типичный уровень доходности большинства клиентов."
- **Выявленные паттерны и инсайты:** Опиши интересные закономерности (например, "наблюдается выраженный рост активности клиентов в четвертом квартале", "категория 'Премиум' доминирует на рынке, составляя 60% продаж").
- **Потенциальные причины и последствия:** Предложи возможные объяснения паттернов и обсуди их потенциальное влияние на бизнес.

==ОБЩИЕ ТРЕНДЫ И ВЗАИМОСВЯЗИ (ВИЗУАЛИЗАЦИИ)**
Обобщи ключевые выводы, сделанные на основе визуализаций. Не просто перечисляй типы графиков, а **расскажи историю, которую они рассказывают**:
- Какие **основные тренды во времени** были обнаружены?
- Какие **взаимосвязи между переменными** выявлены (например, корреляция, различия между группами)?
- Какие **аномалии или выбросы** были замечены?
- Как визуализации **подтверждают или дополняют** выводы из числовых метрик?

==ВЫВОДЫ И СТРАТЕГИЧЕСКИЕ РЕКОМЕНДАЦИИ==
Заверши отчёт **четкими и действенными выводами**.
- **Сводка ключевых инсайтов:** Кратко повтори самые важные пункты из предыдущих разделов.
- **Конкретные рекомендации:** Предложи **практические шаги**, которые можно предпринять на основе анализа. Будь максимально конкретен (например, "Рекомендуется запустить целевую маркетинговую кампанию для клиентов с низкой активностью, чтобы повысить их вовлеченность", "Необходимо провести аудит данных по столбцу 'SUM_FOT' для устранения несогласованности в формате чисел").
- **Направления для дальнейшего исследования:** Предложи идеи для следующих шагов (например, "Рекомендуется углубленный анализ причин высокой доходности у топ-5% клиентов", "Следует построить прогнозную модель для предсказания оттока клиентов").
- **Ограничения анализа:** Честно укажи на ограничения (например, "Анализ основан на данных за один год, что может не отражать долгосрочные тренды", "Наличие пропусков в данных о доходности может повлиять на точность средних значений").

**Форматирование:**
- Используй только простой текст.
- Строго используй указанные заголовки `==ЗАГОЛОВОК==` и подзаголовки `[Имя столбца]`.
- Не используй списки, маркеры, жирный шрифт или другие элементы разметки.
- Пиши на русском языке.
---
"""
# -----------------------------------------
# Настройка приложения
# -----------------------------------------
st.set_page_config(page_title="Анализ данных с LangChain", layout="wide")
st.title("📊 Автоматический анализ данных (CSV/Excel)")

# -----------------------------------------
# Инициализация моделей Ollama
# -----------------------------------------
@st.cache_resource
def get_llms():
    try:
        llm_analyst = Ollama(model="qwen3:8b", temperature=0.55) # Слегка понизил temp для структуры
        llm_coder = Ollama(model="qwen3-coder:latest", temperature=0.2)
        return llm_analyst, llm_coder
    except Exception as e:
        st.error(f"❌ Ошибка подключения к Ollama. Убедитесь, что Ollama запущен и модели загружены. {e}")
        st.stop()

llm_analyst, llm_coder = get_llms()
# Python REPL
repl = PythonREPL()

# -----------------------------------------
# Вспомогательные функции
# -----------------------------------------
def extract_python_code(markdown_text):
    """
    Извлекает Python-код из Markdown-блока, заключенного в ```python ... ```.
    Если блоков несколько, возвращает содержимое первого.
    Если блоков нет, возвращает оригинальный текст.
    """
    if not isinstance(markdown_text, str):
        return str(markdown_text)
    # Регулярное выражение для поиска кода внутри ```python ... ```
    pattern = r"```(?:python|Python)\s*(.*?)\s*```"
    match = re.search(pattern, markdown_text, re.DOTALL)
    if match:
        code = match.group(1)
        return code.strip()
    else:
        # Если блок не найден, удаляем оставшиеся ```
        cleaned_text = re.sub(r"```.*?\n", "", markdown_text)
        cleaned_text = re.sub(r"```", "", cleaned_text)
        return cleaned_text.strip()

def convert_numpy_types(obj):
    """
    Рекурсивно преобразует объекты numpy, pandas и других специфических типов в стандартные типы Python,
    пригодные для сериализации в JSON. ИСКЛЮЧЕНА возможность AttributeError от .item().
    """
    if isinstance(obj, dict):
        new_dict = {}
        for k, v in obj.items():
            if isinstance(k, (np.integer, np.floating)):
                try:
                    new_key = k.item()
                except (AttributeError, ValueError):
                    new_key = str(k)
            elif isinstance(k, (pd.Timestamp, pd.Timedelta)) or hasattr(k, 'isoformat'):
                try:
                    new_key = k.isoformat()
                except Exception:
                    new_key = str(k)
            elif not isinstance(k, (str, int, float, bool)) or k is None:
                new_key = str(k)
            else:
                new_key = k
            new_dict[new_key] = convert_numpy_types(v)
        return new_dict
    if isinstance(obj, list):
        return [convert_numpy_types(item) for item in obj]
    if isinstance(obj, np.integer):
        return int(obj)
    if isinstance(obj, np.floating):
        if np.isnan(obj) or np.isinf(obj):
             return str(obj)
        return float(obj)
    if isinstance(obj, np.bool_):
        return bool(obj)
    if isinstance(obj, np.str_):
        return str(obj)
    if isinstance(obj, np.ndarray):
        try:
            if obj.ndim == 0:
                return convert_numpy_types(obj.item())
            return obj.tolist()
        except Exception as e:
            logger.warning(f"Не удалось преобразовать np.ndarray в list: {e}. Преобразуем в str.")
            return str(obj)
    if isinstance(obj, (pd.Timestamp, pd.Timedelta)):
        try:
            return obj.isoformat()
        except Exception:
            return str(obj)
    if hasattr(obj, 'isoformat'):
        try:
            return obj.isoformat()
        except Exception:
            return str(obj)
    return obj

# Парсинг строкового ответа структуры 
def parse_struct_analyze_response(response_text):
    """
    Парсит строковый ответ от LLM по структуре данных.
    """
    logger.debug(f"Попытка парсинга структуры из строки: {repr(response_text[:300])}...")
    
    parsed_data = {"columns": [], "datetime_candidates": []}
    
    # Извлечение блока столбцов
    columns_match = re.search(r"---COLUMNS_START---(.*?)---COLUMNS_END---", response_text, re.DOTALL)
    if columns_match:
        columns_text = columns_match.group(1).strip()
        # Разбиваем на блоки для каждого столбца
        column_blocks = re.split(r'\n\s*\n', columns_text)
        for block in column_blocks:
            if not block.strip():
                continue
            lines = block.strip().split('\n')
            col_info = {}
            for line in lines:
                if line.startswith("Столбец:"):
                    col_info["name"] = line[len("Столбец:"):].strip()
                elif line.startswith("Тип:"):
                    col_info["type"] = line[len("Тип:"):].strip()
                elif line.startswith("Описание:"):
                    col_info["description"] = line[len("Описание:"):].strip()
            if col_info:
                parsed_data["columns"].append(col_info)
    else:
        logger.warning("Не найден блок ---COLUMNS_START---...---COLUMNS_END---")
        st.warning("⚠️ Не найден блок описания столбцов в ответе LLM.")
        return {}

    # Извлечение блока datetime кандидатов
    datetime_match = re.search(r"---DATETIME_CANDIDATES_START---(.*?)---DATETIME_CANDIDATES_END---", response_text, re.DOTALL)
    if datetime_match:
        datetime_text = datetime_match.group(1).strip()
        if datetime_text:
            parsed_data["datetime_candidates"] = [name.strip() for name in datetime_text.split(',') if name.strip()]
    else:
        logger.warning("Не найден блок ---DATETIME_CANDIDATES_START---...---DATETIME_CANDIDATES_END---")
        # Это не критично, список может быть пустым
        
    logger.info("Структура данных успешно распарсена из строки.")
    logger.debug(f"Распарсенные данные: {parsed_data}")
    return parsed_data


#  Парсинг строкового ответа плана метрик ---
def parse_metrics_plan_response(response_text):
    """
    Парсит строковый ответ от LLM по плану метрик.
    """
    logger.debug(f"Попытка парсинга плана метрик из строки: {repr(response_text[:300])}...")
    
    parsed_data = {}
    
    # Извлечение блока метрик
    metrics_match = re.search(r"---METRICS_START---(.*?)---METRICS_END---", response_text, re.DOTALL)
    if metrics_match:
        metrics_text = metrics_match.group(1).strip()
        # Разбиваем на блоки для каждого столбца
        column_blocks = re.split(r'\n\s*\n', metrics_text)
        for block in column_blocks:
            if not block.strip():
                continue
            lines = block.strip().split('\n')
            col_name = None
            metrics_list = []
            for line in lines:
                if line.startswith("Столбец:"):
                    col_name = line[len("Столбец:"):].strip()
                elif line.startswith("Метрики:"):
                    metrics_str = line[len("Метрики:"):].strip()
                    metrics_list = [m.strip() for m in metrics_str.split(',') if m.strip()]
            
            if col_name:
                parsed_data[col_name] = metrics_list
    else:
        logger.warning("Не найден блок ---METRICS_START---...---METRICS_END---")
        st.warning("⚠️ Не найден блок плана метрик в ответе LLM.")
        return {}

    logger.info("План метрик успешно распарсен из строки.")
    logger.debug(f"Распарсенные данные: {parsed_data}")
    return parsed_data



# safe_code_execution с перезагрузкой df
def static_code_analysis(code_str, context_name=""):
    """
    Выполняет базовую статическую проверку кода.
    Возвращает список предупреждений.
    """
    warnings = []
    try:
        # 1. Проверка синтаксиса с помощью ast
        ast.parse(code_str)
        logger.info(f"Статический анализ ({context_name}): Синтаксис корректен.")
    except SyntaxError as e:
        warnings.append(f"❌ Синтаксическая ошибка: {e}")
        logger.warning(f"Статический анализ ({context_name}): Синтаксическая ошибка: {e}")
        return warnings # Дальнейший анализ бессмысленен при синтаксической ошибке

    # 2. Проверка на потенциально опасные паттерны (упрощенная)
    lines = code_str.split('\n')
    for i, line in enumerate(lines):
        # Проверка на resample без явного указания индекса
        if '.resample(' in line and not ('set_index(' in code_str and code_str.find('set_index(') < code_str.find('.resample(')):
            if line.strip().startswith('df.') or '.resample(' in line.split('=')[-1]:
                 warnings.append(
                    f"⚠️ Потенциальная проблема в строке {i+1}: Использование 'resample' без предшествующего 'set_index' на datetime-столбце может вызвать TypeError. "
                    f"Убедитесь, что индекс является DatetimeIndex или используйте df.set_index('date_column').resample(...). "
                    f"Строка: {line.strip()}"
                )
            logger.warning(f"Статический анализ ({context_name}): Потенциальная проблема resample в строке {i+1}: {line.strip()}")

    if not warnings:
        logger.info(f"Статический анализ ({context_name}): Потенциальные проблемы не обнаружены.")
    else:
        logger.warning(f"Статический анализ ({context_name}): Обнаружены потенциальные проблемы: {warnings}")
    return warnings

def load_df_from_state():
    """Загружает df из st.session_state по сохраненному пути или файлу."""
    if "uploaded_file" in st.session_state and st.session_state["uploaded_file"] is not None:
        # Загрузка из загруженного файла
        uploaded_file = st.session_state["uploaded_file"]
        try:
            if uploaded_file.name.endswith(".csv"):
                try:
                    df = pd.read_csv(uploaded_file, encoding='utf-8')
                except UnicodeDecodeError:
                    try:
                        uploaded_file.seek(0)
                        df = pd.read_csv(uploaded_file, encoding='latin1')
                    except:
                        uploaded_file.seek(0)
                        df = pd.read_csv(uploaded_file, encoding='cp1251')
            elif uploaded_file.name.endswith(".xlsx"):
                df = pd.read_excel(uploaded_file)
            else:
                raise ValueError("Неподдерживаемый формат файла.")
            return df
        except Exception as e:
            st.error(f"❌ Ошибка повторной загрузки загруженного файла: {e}")
            return None
    elif "file_path" in st.session_state and st.session_state["file_path"]:
        # Загрузка по пути
        file_path = st.session_state["file_path"]
        try:
            if file_path.endswith(".csv"):
                try:
                    df = pd.read_csv(file_path, encoding='utf-8')
                except UnicodeDecodeError:
                    try:
                        df = pd.read_csv(file_path, encoding='latin1')
                    except:
                        df = pd.read_csv(file_path, encoding='cp1251')
            elif file_path.endswith(".xlsx"):
                df = pd.read_excel(file_path)
            else:
                raise ValueError("Неподдерживаемый формат файла.")
            return df
        except Exception as e:
            st.error(f"❌ Ошибка повторной загрузки файла по пути {file_path}: {e}")
            return None
    else:
        st.error("❌ Путь к файлу или загруженный файл не найдены в состоянии приложения.")
        return None

# Преобразование дат на основе списка от LLM ---
def preprocess_dates_based_on_llm(df: pd.DataFrame, datetime_columns: list) -> pd.DataFrame:
    """
    Преобразует указанные столбцы в datetime.
    Args:
        df (pd.DataFrame): Исходный DataFrame.
        datetime_columns (list): Список имен столбцов для преобразования.
    Returns:
        pd.DataFrame: DataFrame с преобразованными столбцами.
    """
    df_processed = df.copy()
    successfully_converted = []
    if not datetime_columns:
        logger.info("Список datetime-столбцов пуст. Преобразование не требуется.")
        st.info("LLM не идентифицировала столбцы с датами для преобразования.")
        return df_processed

    logger.info(f"Начало обработки дат для столбцов: {datetime_columns}")
    st.info(f"LLM идентифицировала потенциальные столбцы с датами: {', '.join(datetime_columns)}. Начинается преобразование...")

    for col in datetime_columns:
        if col not in df_processed.columns:
             logger.warning(f"Столбец '{col}', указанный LLM как datetime, не найден в DataFrame.")
             st.warning(f"Столбец '{col}', указанный как datetime, не найден в данных.")
             continue
        try:
            logger.debug(f"Преобразование столбца '{col}'...")
            df_processed[col] = pd.to_datetime(
                df_processed[col],
                # infer_datetime_format=True, # Может быть deprecated
                errors='coerce' # Преобразует недействительные значения в NaT
            )
            successfully_converted.append(col)
            logger.info(f"✅ Столбец '{col}' успешно преобразован в формат datetime64[ns].")
            st.success(f"✅ Столбец '{col}' успешно преобразован в datetime.")
        except Exception as e:
            logger.error(f"❌ Не удалось преобразовать столбец '{col}' в datetime: {e}")
            st.error(f"❌ Не удалось преобразовать столбец '{col}' в datetime: {e}")

    if successfully_converted:
        success_msg = f"Преобразованы следующие столбцы в datetime: {', '.join(successfully_converted)}"
        logger.info(success_msg)
        st.success(success_msg)
    else:
        logger.info("Ни один из указанных столбцов не был успешно преобразован.")
        st.info("Указанные столбцы не были преобразованы из-за ошибок.")

    return df_processed

# Обработка пропусков перед анализом
def handle_missing_values_before_analysis(df: pd.DataFrame, metrics_plan_dict: dict) -> pd.DataFrame:
    """
    Обрабатывает пропущенные значения в DataFrame перед анализом/визуализацией.
    Применяется только к столбцам, упомянутым в metrics_plan_dict.

    Args:
        df (pd.DataFrame): Исходный DataFrame.
        metrics_plan_dict (dict): Словарь {столбец: [метрики]} от LLM.

    Returns:
        pd.DataFrame: DataFrame с обработанными пропусками.
    """
    if not metrics_plan_dict:
        logger.info("План метрик пуст. Обработка пропусков не требуется.")
        return df.copy()

    df_handled = df.copy()
    columns_to_process = list(metrics_plan_dict.keys())
    logger.info(f"Начало обработки пропусков для столбцов: {columns_to_process}")
    st.info(f"Начало обработки пропусков для столбцов, участвующих в анализе: {', '.join(columns_to_process)}")

    for col in columns_to_process:
        if col not in df_handled.columns:
            logger.warning(f"Столбец '{col}' из плана метрик не найден в DataFrame. Пропущен.")
            st.warning(f"Столбец '{col}' из плана метрик не найден в данных. Пропущен.")
            continue

        dtype = df_handled[col].dtype
        logger.debug(f"Обработка столбца '{col}' (тип: {dtype})")

        # 1. Удаление строк с пустыми значениями в datetime столбцах
        if pd.api.types.is_datetime64_any_dtype(dtype):
            initial_rows = len(df_handled)
            df_handled = df_handled.dropna(subset=[col])
            final_rows = len(df_handled)
            logger.info(f"Столбец '{col}' (datetime): удалено {initial_rows - final_rows} строк с NaT.")
            st.info(f"✅ Столбец '{col}' (datetime): удалено {initial_rows - final_rows} строк с пропущенными датами.")
            continue # Переходим к следующему столбцу

        # 2. Заполнение для остальных типов
        if df_handled[col].isna().any():
            if pd.api.types.is_numeric_dtype(dtype):
                # Заполняем числовые столбцы 0
                fill_value = 0
                df_handled[col] = df_handled[col].fillna(fill_value)
                logger.info(f"Столбец '{col}' (числовой): заполнены пропуски значением {fill_value}.")
                st.info(f"✅ Столбец '{col}' (числовой): заполнены пропуски значением {fill_value}.")
            else:
                # Заполняем нечисловые столбцы 'нет данных'
                fill_value = 'нет данных'
                df_handled[col] = df_handled[col].fillna(fill_value)
                logger.info(f"Столбец '{col}' (другой тип): заполнены пропуски значением '{fill_value}'.")
                st.info(f"✅ Столбец '{col}' (другой тип): заполнены пропуски значением '{fill_value}'.")
        else:
            logger.debug(f"Столбец '{col}' не содержит пропусков.")

    logger.info("Обработка пропусков перед анализом завершена.")
    st.success("✅ Обработка пропусков для участвующих в анализе столбцов завершена.")
    return df_handled



def safe_code_execution(code, context_name="", required_imports=None):
    """
    Безопасное выполнение Python кода с обработкой исключений, перепроверкой и перезагрузкой df.
    Всегда добавляет обязательные импорты перед выполнением.
    """
    if required_imports is None:
        required_imports = []
    try:
        logger.info(f"Выполняется {context_name}...")
        # 1. Извлекаем чистый Python-код
        clean_code = extract_python_code(code)
        logger.debug(f"Извлеченный код ({context_name}):\n{textwrap.shorten(clean_code, width=200, placeholder='...')}")

        # 2. Формируем финальный код с обязательными импортами
        final_code_lines = []
        # Добавляем обязательные импорты в начало
        for imp in required_imports:
             final_code_lines.append(imp)
        # Добавляем извлеченный код, разбив его на строки
        final_code_lines.extend(clean_code.split('\n')) # Исправлено: split('\n')
        final_code_to_execute = "\n".join(final_code_lines) # Исправлено: "\n".join()
        logger.debug(f"Финальный код для выполнения ({context_name}):\n{textwrap.shorten(final_code_to_execute, width=300, placeholder='...')}")

        # Статическая перепроверка 
        st.write(f"🔍 Перепроверка сгенерированного кода для **{context_name}**...")
        warnings = static_code_analysis(final_code_to_execute, context_name)
        if warnings:
            st.warning(f"⚠️ Найдены потенциальные проблемы в коде для **{context_name}**:")
            for w in warnings:
                st.markdown(f"* {w}")
            # Логика принятия решения (упрощенная: автоматическое продолжение с предупреждением)
            st.info("ℹ️ Выполнение продолжается, но имейте в виду предупреждения выше...")
        else:
            st.success(f"✅ Код для **{context_name}** прошёл проверку.")


        # Перезагрузка и подготовка df перед выполнением ---
        st.info(f"🔄 Подготовка данных для **{context_name}**...")
        df = load_df_from_state()
        if df is not None:
            # 1. Обработка дат (как и раньше)
            # Получаем список datetime-кандидатов из состояния (если есть)
            datetime_candidates = st.session_state.get("parsed_data_structure", {}).get("datetime_candidates", [])
            if datetime_candidates:
                df = preprocess_dates_based_on_llm(df, datetime_candidates)
            else:
                logger.info("Список datetime-кандидатов не найден. Пропуск обработки дат.")

            # 2. НОВАЯ ЛОГИКА: Обработка пропусков перед выполнением кода анализа/визуализации
            # Проверяем, если это этапы, где это необходимо (расчет метрик или визуализация)
            if context_name in ["расчета метрик", "визуализации"]:
                 # Получаем план метрик из состояния (он должен быть уже создан к этому моменту)
                 metrics_plan_dict = st.session_state.get("metrics_plan_dict", {})
                 if metrics_plan_dict:
                     df = handle_missing_values_before_analysis(df, metrics_plan_dict)
                 else:
                     logger.warning("План метрик не найден в состоянии. Обработка пропусков пропущена.")
                     st.warning("⚠️ План метрик не найден. Обработка пропусков пропущена.")

            # 3. Передаем подготовленный df в среду выполнения
            repl.locals["df"] = df
            logger.info(f"DF успешно подготовлен и передан в repl.locals для {context_name}.")
            st.success(f"✅ Данные подготовлены и загружены для **{context_name}**.")
        else:
            st.error(f"❌ Не удалось перезагрузить данные для **{context_name}**. Выполнение пропущено.")
            return f"Ошибка: Не удалось перезагрузить данные для {context_name}."
        #

        # 3. Выполняем код
        output = repl.run(final_code_to_execute)
        logger.info(f"{context_name} выполнен успешно.")
        return output.strip()
    except Exception as e:
        error_msg = f"❌ Ошибка при выполнении {context_name}: {str(e)}"
        logger.error(error_msg)
        st.error(error_msg)
        # Показываем финальный код, который вызвал ошибку
        if 'final_code_to_execute' in locals():
            with st.expander("Код, вызвавший ошибку"):
                st.code(final_code_to_execute, language="python")
        else:
            with st.expander("Код, вызвавший ошибку"):
                st.code(code, language="python")
        return f"Ошибка выполнения: {e}"


def get_df_info(df, title="DataFrame"):
    """Получает строковое представление информации о DataFrame для отладки."""
    if df is None:
        return f"{title} is None"
    if df.empty:
        return f"{title} is empty"
    buffer = StringIO()
    df.info(buf=buffer)
    info_str = buffer.getvalue()
    return (
        f"--- {title} ---\n"
        f"Shape: {df.shape}\n"
        f"Columns: {list(df.columns)}\n"
        f"Dtypes:\n{df.dtypes.to_string()}\n"
        f"Head:\n{df.head(2).to_string()}\n"
        f"Info:\n{info_str}\n"
        f"-------------------\n"
    )


# -----------------------------------------
# Ввод: путь к файлу
# -----------------------------------------
st.header("1. Укажите путь к файлу")

# Инициализация переменных состояния
if "file_path" not in st.session_state:
    st.session_state["file_path"] = ""

df = None
file_path = st.text_input("Введите полный путь к файлу (CSV или Excel):", value=st.session_state["file_path"])

# Сохраняем путь в состоянии
st.session_state["file_path"] = file_path

if not file_path:
    st.info("Введите путь к файлу.")
    st.stop()

if not os.path.exists(file_path):
    st.error("❌ Файл не найден. Проверьте путь.")
    st.stop()

try:
    if file_path.endswith(".csv"):
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            try:
                df = pd.read_csv(file_path, encoding='latin1')
            except:
                df = pd.read_csv(file_path, encoding='cp1251')
    elif file_path.endswith(".xlsx"):
        df = pd.read_excel(file_path)
    else:
        st.error("❌ Поддерживаются только .csv и .xlsx файлы.")
        st.stop()
except Exception as e:
    st.error(f"❌ Ошибка чтения файла по пути: {e}")
    st.exception(e)
    st.stop()

if df is None or df.empty:
    st.error("❌ Не удалось загрузить данные или файл пуст.")
    st.stop()

st.success("✅ Файл успешно загружен!")
st.write("### 📋 Первые строки данных (исходные):")
st.dataframe(df.head())
logger.info("DF после первоначальной загрузки:\n" + get_df_info(df, "df после загрузки"))

# -----------------------------------------
# Ввод: путь для сохранения результатов
# -----------------------------------------
st.header("2. Укажите путь для сохранения результатов")
output_dir = st.text_input("Путь к директории для сохранения графиков и отчёта:", value=".")
if not output_dir:
    st.warning("Укажите директорию для сохранения.")
    st.stop()
try:
    os.makedirs(output_dir, exist_ok=True)
    st.success(f"✅ Директория готова: {output_dir}")
except Exception as e:
    st.error(f"❌ Не удалось создать/проверить директорию: {e}")
    st.stop()

test_file = os.path.join(output_dir, ".test_write")
try:
    with open(test_file, "w") as f:
        f.write("test")
    os.remove(test_file)
except Exception as e:
    st.error(f"❌ Нет прав на запись в указанную директорию: {e}")
    st.stop()

st.session_state["output_dir"] = output_dir

# -----------------------------------------
# Готовим цепочку LangChain
# -----------------------------------------
st.header("3. Запуск анализа")
if st.button("🚀 Запустить анализ"):
    with st.spinner("Выполняется анализ..."):
        # Загружаем df для получения информации о структуре
        df_for_info = load_df_from_state()
        if df_for_info is None:
            st.error("❌ Не удалось загрузить данные для получения информации о структуре.")
            st.stop()

        # --- Шаг 0: Получение базовой информации о df для промпта LLM ---
        logger.info("DF перед началом анализа:\n" + get_df_info(df_for_info, "df перед началом анализа"))
        buffer = StringIO()
        df_for_info.info(buf=buffer)
        df_info = buffer.getvalue()
        file_info_summary = (
            f"DataFrame Info (сырые данные):\n{df_info}\n"
            f"DataFrame Head:\n{df_for_info.head(10).to_string()}\n"
            f"DataFrame Dtypes:\n{df_for_info.dtypes.to_string()}\n"
            f"DataFrame Shape: {df_for_info.shape}\n"
        )
        logger.info("Подготовлена сводная информация о сырых данных.")

        # --- Шаг 1: Анализ структуры с помощью LLM ---
        with st.spinner("🧠 Анализ структуры данных LLM..."):
            prompt_structure = PromptTemplate.from_template(struct_analyze)
            chain_structure = LLMChain(llm=llm_analyst, prompt=prompt_structure, output_key="data_structure")
            try:
                # Получаем "сырой" ответ 
                result_structure_raw = chain_structure.run(file_info=file_info_summary)
                logger.info(f"Ответ LLM (анализ структуры):\n{result_structure_raw}")

                #  Используем новый парсер строк 
                parsed_structure = parse_struct_analyze_response(result_structure_raw)
                
                
                if not parsed_structure:
                    st.error("❌ Не удалось распарсить строковый ответ от LLM по структуре данных.")
                    # Отображаем необработанный ответ для отладки
                    with st.expander("Необработанный ответ LLM (Структура)"):
                         st.code(result_structure_raw, language="text")
                    st.stop() # Останавливаем выполнение, если структура не получена

                # --- Сохраняем как сырые данные, так и распарсенный словарь ---
                st.session_state["data_structure_raw"] = result_structure_raw # Сохраняем для отладки
                st.session_state["data_structure"] = json.dumps(parsed_structure, indent=2, ensure_ascii=False) # Для отображения
                st.session_state["parsed_data_structure"] = parsed_structure # Сохраняем распарсенный словарь для дальнейшего использования
                logger.info("Структура данных проанализирована LLM.")
                st.success("✅ Структура данных проанализирована LLM.")

            except Exception as e:
                st.error(f"❌ Ошибка при анализе структуры LLM: {e}")
                # Отображаем необработанный ответ даже при других ошибках
                if 'result_structure_raw' in locals():
                    with st.expander("Необработанный ответ LLM (Структура)"):
                        st.code(result_structure_raw, language="text")
                st.stop()

        # --- Шаг 2: Метрики (используем обработанный df для получения актуальной структуры) ---
        with st.spinner("📏 Генерация плана метрик..."):
            # Обновляем информацию о структуре df для промпта метрик
            # Загружаем df заново и применяем преобразование дат
            df_for_processing = load_df_from_state()
            if df_for_processing is None:
                st.error("❌ Не удалось загрузить данные для подготовки к генерации плана метрик.")
                st.stop()
            
            # Применяем преобразование типов дат на основе анализа LLM
            datetime_candidates = st.session_state.get("parsed_data_structure", {}).get("datetime_candidates", [])
            df_processed = preprocess_dates_based_on_llm(df_for_processing, datetime_candidates)
            
            buffer_processed = StringIO()
            df_processed.info(buf=buffer_processed)
            df_info_processed = buffer_processed.getvalue()
            file_info_summary_processed = (
                f"DataFrame Info (после обработки дат):\n{df_info_processed}\n"
                f"DataFrame Head:\n{df_processed.head(10).to_string()}\n"
                f"DataFrame Dtypes:\n{df_processed.dtypes.to_string()}\n"
                f"DataFrame Shape: {df_processed.shape}\n"
            )

            #  Ужесточенный промпт 
            prompt_metrics = PromptTemplate.from_template(m_plan) 
            # 
            chain_metrics_plan = LLMChain(llm=llm_analyst, prompt=prompt_metrics, output_key="metrics_plan")
            try:
                # Передаем обновленную информацию о структуре
                result_metrics_plan_raw = chain_metrics_plan.run(data_structure=file_info_summary_processed)
                logger.info(f"Ответ LLM (план метрик):\n{result_metrics_plan_raw}") # --- ДОБАВЛЕНИЕ: Логируем ответ ---
                
                # Используем новый парсер строк 
                metrics_plan_dict = parse_metrics_plan_response(result_metrics_plan_raw)

                
                if not metrics_plan_dict: # --- ДОБАВЛЕНИЕ: Явная проверка ---
                    st.error("❌ Не удалось распарсить строковый ответ от LLM по плану метрик.")
                    # Отображаем необработанный ответ для отладки
                    with st.expander("Необработанный ответ LLM (План метрик)"):
                         st.code(result_metrics_plan_raw, language="text")
                    st.stop() # Останавливаем выполнение
            
                
                st.session_state["metrics_plan"] = json.dumps(metrics_plan_dict, indent=2, ensure_ascii=False)
                st.session_state["metrics_plan_dict"] = metrics_plan_dict
                logger.info("План метрик сгенерирован.")
                st.success("✅ План метрик сгенерирован.")
            except Exception as e: # Показываем ответ даже при других ошибках 
                st.error(f"❌ Ошибка при генерации плана метрик: {e}")
                if 'result_metrics_plan_raw' in locals():
                    with st.expander("Необработанный ответ LLM (План метрик)"):
                        st.code(result_metrics_plan_raw, language="text")
                st.session_state["metrics_plan"] = "{}"
                st.session_state["metrics_plan_dict"] = {}
                st.stop()
         

        # --- Шаг 3: Генерация кода расчёта ---
        with st.spinner("💻 Генерация кода для расчёта метрик..."):
            # Обновляем информацию о структуре df для промпта генерации кода
            df_structure_info = f"DataFrame имеет {df_processed.shape[0]} строк и {df_processed.shape[1]} столбцов.\n"
            df_structure_info += "Типы данных столбцов (после обработки):\n"
            for col in df_processed.columns:
                df_structure_info += f"  - {col}: {df_processed[col].dtype}\n"


            # Запрашиваем вывод Python-словаря, а не JSON
            prompt_code_gen = PromptTemplate.from_template(
                """Ты — эксперт по Python и анализу данных. Твоя задача — написать корректный и эффективный Python-код для расчёта заданных метрик по указанным столбцам DataFrame с выводом через print(metrics_results).

У тебя есть DataFrame `df`, уже загруженный в среду. Структура df:
{df_structure_info}

Необходимо рассчитать следующие метрики для соответствующих столбцов:
{metrics_plan}

Следуй этим правилам:

- все вычисления производи в глобальносм уровне
- Используй только pandas, numpy. Не добавляй строки импорта — они уже доступны.
- Рассчитывай ТОЛЬКО те метрики, которые указаны в `{metrics_plan}` для каждого столбца.
- ПРЕЖДЕ чем обращаться к `df["имя_столбца"]`, ОБЯЗАТЕЛЬНО проверяй, существует ли столбец в `df.columns`. Если нет — пропусти.
- Для каждого столбца:
  - Создай копию: `series = df[col].copy()`
  - Все пропущенные значения (NaN, NaT) должны быть обработаны как `None` в результате. Не оставляй `nan`, `nat` или `np.nan`.
  #
- Для метрик, требующих resample (например, dates_per_month, dates_per_year):
1. Не используй series.resample(...) напрямую, если series.index не является DatetimeIndex.
2. Вместо этого создай временную серию, где индекс - это сами значения datetime
#

- Преобразование типов:
  - Все скалярные значения (результаты метрик) должны быть в типах Python: `int`, `float`, `str`, `bool`, `None`.
  - Используй: `int(x)` для целых, `float(x)` для вещественных, `str(x)` для дат, `.tolist()` для списков.
  - Если значение — `NaN`, `NaT` или `None`, результат должен быть `None`.
  - Для числовых метрик (mean, std и т.д.) используй только колонки с типом `int64`, `float64`, или те, которые можно безопасно преобразовать в числовой тип через `pd.to_numeric(series, errors='coerce')`. Если после конвертации все значения — NaN, считай колонку нечисловой.

- Обработка пропусков:
  - Не изменяй исходные данные.
  - Вместо `series.fillna(None)` используй: `series = series.where(pd.notna(series), None)` — это безопасно для всех типов.
  - Или, если метрика требует удаления пропусков (например, mean), используй `series.dropna()` **только для расчёта этой метрики**, не изменяя `series` глобально.

- Для datetime-метрик:
  - Убедись, что `series.dtype == 'datetime64[ns]'`.
  - Для `dates_per_month`: создай временную серию: `temp = pd.Series(1, index=pd.to_datetime(series.dropna()))`, затем `temp.resample('M').count().tolist()`.

- При написании кода с NumPy-массивами используй явные условия: np.any(condition) или np.all(condition) вместо if condition:
- Для `mad`: используй формулу: `mad_val = (series - series.mean()).abs().mean()`, если `series` числовая.
- Не изменяй `df`, не сбрасывай индекс, не предполагай тип индекса.
- Избегай дублирования вычислений (например, не считай `quantile(0.25)` дважды).
- Код должен быть свободен от синтаксических ошибок и готов к выполнению.
- Выведи результат одной строкой: print(metrics_results)

Верни только Python-код без пояснений (выполненный в глобальном уровне), комментариев или дополнительного текста."""
#                 """
#             Ты — эксперт по Python и анализу данных. Твоя задача — написать корректный и эффективный Python-код для расчёта заданных метрик по указанным столбцам DataFrame с выводом при помощи print(metrics_results)

# У тебя есть DataFrame `df`, уже загруженный в среду. Структура df:
# {df_structure_info}

# Необходимо рассчитать следующие метрики для соответствующих столбцов:
# {metrics_plan}

# Следуй этим правилам:
# - Используй только pandas, numpy и scipy (если нужно). Не добавляй строки импорта — они уже доступны.
# - Рассчитывай ТОЛЬКО те метрики, которые указаны в `{metrics_plan}` для каждого столбца. 
# - ПРЕЖДЕ ЧЕМ обращаться к `df["имя_столбца"]`, ОБЯЗАТЕЛЬНО проверяй, существует ли такой столбец в `df.columns`
# - Перед расчетом метрик для числового столбца убедись, что он имеет числовой тип (int, float). 
# - Преобразуй все значения из типов NumPy (например, numpy.int64, numpy.float64) в стандартные типы Python: используй `int()`, `float()`, `bool()` для скаляров, `.tolist()` для массивов, `str()` для дат.
# - Обрабатывай пропущенные значения (NaN, NaT) — заменяй их на `None`.
# - Выведи результат ТОЛЬКО ОДНОЙ строкой в формате Python-словаря: print(metrics_results)

# - Не используй вспомогательные функции, не дублируй код, не создавай лишние переменные.
# - Не изменяй исходный DataFrame (`df`), не сбрасывай индекс, не предполагай структуру индекса.
# - Убедись, что код не содержит синтаксических ошибок и готов к выполнению.
# - Для метрик, требующих resample (например, dates_per_month, dates_per_year):
# 1. Не используй series.resample(...) напрямую, если series.index не является DatetimeIndex.
# 2. Вместо этого создай временную серию, где индекс - это сами значения datetime

# -Для расчета `mad` (Mean Absolute Deviation) НЕ используй `series.mad()` напрямую, если не уверен в его существовании. Вместо этого используй эквивалентную формулу: `mad_val = series.sub(series.mean()).abs().mean()`.

# Верни только Python-код без пояснений, комментариев или дополнительного текста. Последняя строка — вывод структурированного текста с результатами.
# """
            )


            chain_code_gen = LLMChain(llm=llm_coder, prompt=prompt_code_gen, output_key="calculation_code")
            metrics_plan_for_prompt = st.session_state.get("metrics_plan", "{}")
            if "metrics_plan_dict" in st.session_state and st.session_state["metrics_plan_dict"]:
                 metrics_plan_for_prompt = json.dumps(st.session_state["metrics_plan_dict"], indent=2, ensure_ascii=False)
            try:
                result_code_gen = chain_code_gen.run(metrics_plan=metrics_plan_for_prompt, df_structure_info=df_structure_info)
                st.session_state["calculation_code"] = result_code_gen
                logger.info("Код для расчёта метрик сгенерирован.")
                st.success("✅ Код для расчёта метрик сгенерирован.")
            except Exception as e:
                st.error(f"❌ Ошибка при генерации кода расчёта метрик: {e}")
                st.session_state["calculation_code"] = "# Код не был сгенерирован из-за ошибки."
                st.stop()

        # --- Шаг 4: Выполнение сгенерированного кода для расчета метрик ---
        # --- Шаг 4: Выполнение сгенерированного кода для расчета метрик ---
        with st.spinner("📊 Выполняется расчет метрик..."):
            required_imports_for_metrics = [
                "import pandas as pd",
                "import numpy as np",
                # json и ast не требуются, если мы не парсим
            ]
            # Выполняем код и получаем вывод как строку
            metrics_results_raw_output = safe_code_execution(
                st.session_state["calculation_code"],
                "расчета метрик",
                required_imports=required_imports_for_metrics
            )
            
            # Сохраняем "сырой" вывод как строку, без попыток парсинга
            st.session_state["metrics_results_raw"] = metrics_results_raw_output
            
            # Поскольку парсинг убран, инициализируем metrics_results как пустой словарь
            # или как саму строку, если ваша логика где-то еще полагается на него.
            # Если нигде больше не используется как словарь, можно и не создавать.
            # st.session_state["metrics_results"] = {} # Опционально, если где-то ожидается

            logger.info("Метрики рассчитаны (вывод сохранен как строка).")
            st.success("✅ Метрики рассчитаны! Результат сохранен как строка.")

            # --- НОВОЕ: Сохранение сгенерированного кода расчета метрик в .py ---
            try:
                # Получаем директорию для сохранения из состояния
                output_dir_for_code = st.session_state.get("output_dir", ".")
                calculation_code_filename = "generated_calculation_code.py"
                calculation_code_path = os.path.join(output_dir_for_code, calculation_code_filename)

                # Получаем содержимое кода
                calculation_code_content = st.session_state.get("calculation_code", "# Код не был сгенерирован.")

                # Убедимся, что содержимое - это строка
                if not isinstance(calculation_code_content, str):
                    calculation_code_content = str(calculation_code_content)

                # Записываем код в файл .py
                with open(calculation_code_path, "w", encoding="utf-8") as f_py:
                    # Можно добавить комментарий в начало файла
                    f_py.write("# Сгенерированный код для расчета метрик\n")
                    f_py.write("# -------------------------------------\n\n")
                    f_py.write(calculation_code_content)

                logger.info(f"Сгенерированный код расчета метрик сохранен в PY: {calculation_code_path}")
                # Опционально: вывести сообщение об успехе
                st.success(f"✅ Сгенерированный код расчета метрик сохранен в PY: {calculation_code_path}")
            except Exception as e_code_save:
                error_msg_code_save = f"❌ Ошибка при сохранении сгенерированного кода расчета метрик в PY: {e_code_save}"
                logger.error(error_msg_code_save)
                st.error(error_msg_code_save)
            # (Весь старый try...except с json.loads или ast.literal_eval здесь был)
            
            # --- Проверка на наличие ошибок в выводе (опционально, но полезно) ---
            if isinstance(metrics_results_raw_output, str) and ("Ошибка выполнения" in metrics_results_raw_output or "Traceback" in metrics_results_raw_output):
                 st.error("❌ Похоже, при выполнении кода расчета метрик произошла ошибка.")
                 st.code(metrics_results_raw_output, language="text")
                 st.stop() # Останавливаем, если код не выполнился успешно
            # --- Конец проверки ---

        # --- Шаг 5: Анализ результатов и сохранение отчета ---
        with st.spinner("🔍 Анализ рассчитанных метрик..."):
            prompt_analysis = PromptTemplate.from_template(data_analyze)
            # Передаем оригинальный "сырой" вывод в промпт анализа
            prompt_analysis = prompt_analysis.partial(metrics_results_raw=st.session_state["metrics_results_raw"])
            chain_analysis = LLMChain(llm=llm_analyst, prompt=prompt_analysis, output_key="analysis_summary")
            try:
                result_analysis = chain_analysis.invoke({})
                raw_analysis_summary = result_analysis.get("analysis_summary", "Анализ не выполнен.")
                st.session_state["analysis_summary"] = convert_numpy_types(raw_analysis_summary)
                logger.info("Анализ метрик завершён.")
                st.success("✅ Анализ метрик завершён.")

                # --- НОВОЕ: Сохранение отчета об анализе метрик ---
                # Получаем директорию для сохранения из состояния
                output_dir_for_analysis = st.session_state.get("output_dir", ".")
                analysis_report_filename_base = "analysis_summary_report"
                analysis_report_txt_path = os.path.join(output_dir_for_analysis, f"{analysis_report_filename_base}.txt")
                analysis_report_docx_path = os.path.join(output_dir_for_analysis, f"{analysis_report_filename_base}.docx")

                # Убедимся, что содержимое для сохранения - это строка
                analysis_content_to_save = st.session_state["analysis_summary"]
                if not isinstance(analysis_content_to_save, str):
                    analysis_content_to_save = str(convert_numpy_types(analysis_content_to_save))

                # Сохранение в .txt
                try:
                    with open(analysis_report_txt_path, "w", encoding="utf-8") as f_txt:
                        f_txt.write(analysis_content_to_save)
                    logger.info(f"Отчет об анализе метрик сохранен в TXT: {analysis_report_txt_path}")
                    st.success(f"✅ Отчет об анализе метрик сохранен в TXT: {analysis_report_txt_path}")
                except Exception as e_txt:
                    error_msg_txt = f"❌ Ошибка при сохранении отчета об анализе метрик в TXT: {e_txt}"
                    logger.error(error_msg_txt)
                    st.error(error_msg_txt)

                # Сохранение в .docx
                try:
                    from docx import Document # Импортируем внутри блока try, чтобы не ломать всё при отсутствии библиотеки
                    doc = Document()
                    # Добавляем заголовок
                    doc.add_heading('Анализ рассчитанных метрик', level=1)
                    # Добавляем содержимое отчета как параграф(ы)
                    # splitlines и добавление построчно может помочь с форматированием
                    for line in analysis_content_to_save.splitlines():
                        if line.strip(): # Добавляем только непустые строки как отдельные параграфы
                            doc.add_paragraph(line)
                        # Если весь текст нужно в один параграф, используйте:
                        # doc.add_paragraph(analysis_content_to_save)
                    doc.save(analysis_report_docx_path)
                    logger.info(f"Отчет об анализе метрик сохранен в DOCX: {analysis_report_docx_path}")
                    st.success(f"✅ Отчет об анализе метрик сохранен в DOCX: {analysis_report_docx_path}")
                except ImportError as e_import:
                    error_msg_import = f"❌ Библиотека python-docx не установлена. Отчет в DOCX не сохранен: {e_import}"
                    logger.error(error_msg_import)
                    st.error(error_msg_import)
                except Exception as e_docx:
                    error_msg_docx = f"❌ Ошибка при сохранении отчета об анализе метрик в DOCX: {e_docx}"
                    logger.error(error_msg_docx)
                    st.error(error_msg_docx)
                # --- КОНЕЦ НОВОГО ---

            except Exception as e:
                st.error(f"❌ Ошибка при анализе метрик: {e}")
                st.session_state["analysis_summary"] = "Ошибка анализа."

        # --- Шаг 6: Генерация кода визуализации ---
        with st.spinner("🎨 Генерация кода визуализации..."):
            
            prompt_viz = PromptTemplate.from_template(
               """# 🎯 ЦЕЛЬ
Построй графики, иллюстрирующие ОСНОВНЫЕ закономерности из анализа метрик. Каждый график должен показывать УНИКАЛЬНУЙ инсайт. НЕ ДУБЛИРУЙ ИДЕИ.

# ⚠️ ОГРАНИЧЕНИЯ
- Построй ровно 20 графиков, не больше, не меньше

# 📊 ТИПЫ ГРАФИКОВ (рекомендуемое распределение)
- Распределения числовых переменных (hist, kde, box, violin)
- Связь числовой и categorиальной (boxplot/violinplot)
- Связь двух числовых (scatter с hue, regplot)
- Категориальные (barplot, countplot)
- Временные ряды (агрегация по периодам)
- Взаимосвязи (heatmap корреляций, crosstab)
- Уникальные инсайты (аномалии, сравнения групп)

# 🎨 ЦВЕТОВЫЕ ТРЕБОВАНИЯ
- Все графики должны быть цветными.
- В barplot и аналогичных графиках КАЖДЫЙ столбец должен иметь свой цвет (используй palette, hue, или color для каждого столбца)
- Используй палитры: `palette='tab10'`, `cmap='viridis'`, `color=...`, `hue=...`.
- Применяй `hue`, `palette`, `cmap` везде, где возможно.
- Не оставляй графики черно-белыми или однотонными.
- В barplot используй `palette` или передавай список цветов, чтобы каждый столбец был разного цвета.

# 🚫 ЗАПРЕЩЕНО
- Дублирование идей (например, несколько графиков для одного столбца без новой информации)
- Графики без инсайтов
- Пропуск графиков из-за большого количества категорий — используй top-N
- Более 20 графиков
- Однотонные barplot (все столбцы одного цвета)

# 📥 Входные данные
У тебя есть:
- DataFrame `df`
- Структура: {df_structure_info}
- Метрики: {metrics_results_raw}
- Анализ: {analysis_summary}

# ✅ ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ
- Используй `matplotlib`, `seaborn`
- Сохраняй в `{output_dir}` как `plot_{{столбцы}}_{{тип}}.png`
- Не используй `plt.show()`, только `plt.savefig()`
- Для каждого графика используй `try-except`
- Используй `palette='tab10'`, `cmap='viridis'`
- Для оси X с >5 меток: `rotation=45`, `ha='right'`
- Всегда `plt.tight_layout()`
- Создай папку: `os.makedirs('{output_dir}', exist_ok=True)`
- Не изменяй `df`, не загружай заново
- Верни ТОЛЬКО код (20 графиков!!! с комментариями вроде #График 1: ... ). 
"""
           )
           
            # df_structure_info уже определена выше
            # prompt_viz = PromptTemplate.from_template(
            #     """
            #     Ты получил DataFrame `df`, который уже загружен в среду выполнения Python. Структура df:
            #     {df_structure_info}
            #     Также ты получил следующие рассчитанные метрики и анализ:
            #     --- Метрики ---
            #     {metrics_results_raw}
            #     --- Анализ ---
            #     {analysis_summary}
            #     Твоя задача: сгенерировать Python-код для визуализации данных из `df` с использованием matplotlib и seaborn.
            #     Требования:
            #     - Для числовых столбцов: гистограммы, boxplot, KDE.
            #     - Для категориальных: bar plot с частотами.
            #     - Для дат (datetime): линейные графики временных рядов, гистограммы по периодам (месяцам, годам).
            #     - Если более одного числового столбца: heatmap корреляций.
            #     - Все графики сохраняй в отдельные файлы в папку '{output_dir}'.
            #     - Имя файла: plot_{{столбец}}_{{тип}}.png
            #     - Не показывай графики (plt.show()), только сохраняй (plt.savefig).
            #     - Обрабатывай возможные ошибки (например, пустые данные, нечисловые значения).
            #     - Учитывай, что некоторые столбцы могут быть типа `datetime64[ns]`.
            #     - Используй информацию о структуре df выше, чтобы выбрать правильные столбцы для правильных типов графиков.
            #     - НЕ пытайся загружать данные снова (например, с помощью pd.read_csv). Используй переменную `df`, которая уже существует.
            #     - НЕ переопределяй переменную `df`.
            #     - НЕ используй `del df`.
            #     - НЕ используй `df.resample(...)` напрямую на df или столбце, если индекс не является DatetimeIndex. Сначала установи datetime-столбец как индекс: `df.set_index('date_column')`.
            #     - НЕ пиши ```python
            #     Включи в код создание папки, если её нет.
            #     Верни ТОЛЬКО код. Без пояснений. Убедись, что код импортирует все нужные библиотеки.
            #     """
            # )
            
#            prompt_viz = PromptTemplate.from_template(
# """
# # ⚠️ ВАЖНО: СТРОГОЕ ОГРАНИЧЕНИЕ
# Построй РОВНО 20 ГРАФИКОВ. Ни больше, ни меньше. Если ты построил 20 — код завершается. Не добавляй больше.

# # 🎯 ЦЕЛЬ
# Каждый график должен иллюстрировать УНИКАЛЬНУЮ закономерность из анализа. НЕ ДУБЛИРУЙ ИДЕИ.

# # 📊 РАСПРЕДЕЛЕНИЕ (обязательно)
# - 2: распределение числовых переменных (разные столбцы, разные типы: hist, kde, box, violin)
# - 3: связь числовой и категориальной (boxplot/violinplot для пар с наибольшими различиями)
# - 4: связь двух числовых (scatter с hue, regplot)
# - 5: категориальные (barplot, top-15)
# - 2: временные ряды (агрегация по месяцам)
# - 2: взаимосвязи (heatmap корреляций, heatmap crosstab)
# - 2: уникальные инсайты (например, аномалии, сравнение групп)

# # 🎨 ЦВЕТОВЫЕ ТРЕБОВАНИЯ
# - Все графики должны быть цветными.
# - Используй палитры: `palette='tab10'`, `cmap='viridis'`, `color=...`, `hue=...`.
# - Для scatterplot, barplot, boxplot и т.д. применяй `hue`, `palette`, `cmap` везде, где возможно.
# - Не оставляй графики черно-белыми или однотонными.

# # 🚫 ЗАПРЕЩЕНО
# - >20 графиков
# - Дублирование (например, hist + kde + box для одного столбца)
# - Графики для столбцов без инсайтов
# - Пропуск графиков из-за большого количества категорий — используй top-N

# # 📥 Входные данные
# У тебя есть:
# - DataFrame `df`
# - Структура: {df_structure_info}
# - Метрики: {metrics_results_raw}
# - Анализ: {analysis_summary}

# # ✅ Требования к коду
# - Используй `matplotlib`, `seaborn`
# - Сохраняй в `{output_dir}` как `plot_{{столбцы}}_{{тип}}.png`
# - Не используй `plt.show()`, только `plt.savefig()`
# - Каждый график в `try-except`
# - Используй `palette='tab10'`, `cmap='viridis'`
# - Для оси X с >5 меток: `rotation=45`, `ha='right'`
# - Всегда `plt.tight_layout()`
# - Создай папку: `os.makedirs('{output_dir}', exist_ok=True)`
# - Не изменяй `df`, не загружай заново
# - Верни ТОЛЬКО код. Без пояснений. Без ```python
# """)

      
        prompt_viz = prompt_viz.partial(
            df_structure_info=df_structure_info,
            metrics_results_raw=st.session_state["metrics_results_raw"],
            analysis_summary=st.session_state["analysis_summary"],
            output_dir=output_dir
        )
        chain_viz_code = LLMChain(llm=llm_coder, prompt=prompt_viz, output_key="viz_code")
        try:
            result_viz_code = chain_viz_code.invoke({})
            raw_viz_code = result_viz_code.get("viz_code", "# Код визуализации не сгенерирован.")
            if not isinstance(raw_viz_code, str):
                st.session_state["viz_code"] = str(convert_numpy_types(raw_viz_code))
            else:
                st.session_state["viz_code"] = raw_viz_code
            logger.info("Код визуализации сгенерирован.")
            st.success("✅ Код визуализации сгенерирован.")
        except Exception as e:
            st.error(f"❌ Ошибка при генерации кода визуализации: {e}")
            st.session_state["viz_code"] = f"# Ошибка генерации: {e}\n# Код визуализации не сгенерирован из-за ошибки."

        # --- Шаг 7: Выполнение кода визуализации ---
        with st.spinner("🖼️ Выполняется построение графиков..."):
            required_imports_for_viz = [
                "import pandas as pd",
                "import numpy as np",
                "import matplotlib",
                "matplotlib.use('Agg')",
                "import matplotlib.pyplot as plt",
                "import seaborn as sns",
                "import json",
                "import os"
            ]
            viz_output = safe_code_execution(
                st.session_state["viz_code"],
                "визуализации",
                required_imports=required_imports_for_viz
            )
            if "error" not in viz_output.lower() and "отменено" not in viz_output.lower():
                st.success("✅ Графики сохранены.")
            logger.info("Визуализация завершена.")
        try:
                        # Получаем директорию для сохранения из состояния
            output_dir_for_viz_code = st.session_state.get("output_dir", ".")
            viz_code_filename = "generated_visualization_code.py"
            viz_code_path = os.path.join(output_dir_for_viz_code, viz_code_filename)

            # Получаем содержимое кода
            viz_code_content = st.session_state.get("viz_code", "# Код визуализации не был сгенерирован.")

            # Убедимся, что содержимое - это строка
            if not isinstance(viz_code_content, str):
                viz_code_content = str(viz_code_content)

            # Записываем код в файл .py
            with open(viz_code_path, "w", encoding="utf-8") as f_py:
                # Можно добавить комментарий в начало файла
                f_py.write("# Сгенерированный код для визуализации\n")
                f_py.write("# ----------------------------------\n\n")
                f_py.write(viz_code_content)

            logger.info(f"Сгенерированный код визуализации сохранен в PY: {viz_code_path}")
            # Опционально: вывести сообщение об успехе
            st.success(f"✅ Сгенерированный код визуализации сохранен в PY: {viz_code_path}")
        except Exception as e_viz_code_save:
            error_msg_viz_code_save = f"❌ Ошибка при сохранении сгенерированного кода визуализации в PY: {e_viz_code_save}"
            logger.error(error_msg_viz_code_save)
            st.error(error_msg_viz_code_save)

        # --- Шаг 8: Генерация итогового отчёта ---
        with st.spinner("📄 Генерация итогового отчёта..."):
            # Проверка наличия необходимых данных для генерации отчета
            if "analysis_summary" not in st.session_state or not st.session_state["analysis_summary"]:
                error_msg = "❌ Невозможно сгенерировать итоговый отчёт: анализ метрик не был выполнен или отсутствует."
                st.error(error_msg)
                st.session_state["final_report"] = error_msg # Записываем ошибку в состояние
                logger.error("Попытка генерации итогового отчёта без анализа (analysis_summary отсутствует в st.session_state).")
                st.stop() # Останавливаем выполнение

            prompt_report = PromptTemplate.from_template(final_rep)
            # Добавляем partial для передачи analysis_summary в промпт
            prompt_report = prompt_report.partial(analysis_summary=st.session_state["analysis_summary"])
            chain_report = LLMChain(llm=llm_analyst, prompt=prompt_report, output_key="final_report")
            try:
                # --- ИЗМЕНЕНИЕ: Сначала выполняем цепочку, чтобы получить результат ---
                # Используем invoke, как в других местах, и передаем пустой словарь
                result_report = chain_report.invoke({}) # Выполняем генерацию
                # Получаем сгенерированный отчет из результата
                raw_final_report = result_report.get("final_report", "Итоговый отчет не сгенерирован.")
                # Конвертируем numpy типы, если необходимо
                final_report_content = convert_numpy_types(raw_final_report)
                # Сохраняем сгенерированный отчет в st.session_state
                st.session_state["final_report"] = final_report_content
                # --- КОНЕЦ ИЗМЕНЕНИЯ ---

                # --- НОВОЕ: Теперь, когда отчет есть в состоянии, сохраняем его в файлы ---
                # Получаем директорию для сохранения из состояния
                output_dir = st.session_state.get("output_dir", ".") # Убедитесь, что output_dir уже установлен
                final_report_filename_base = "final_report"
                final_report_txt_path = os.path.join(output_dir, f"{final_report_filename_base}.txt")
                final_report_docx_path = os.path.join(output_dir, f"{final_report_filename_base}.docx")

                # Убедимся, что содержимое для сохранения - это строка
                # final_report_content уже конвертирован выше
                if not isinstance(final_report_content, str):
                    final_report_content_to_save = str(final_report_content)
                else:
                    final_report_content_to_save = final_report_content

                # Сохранение в .txt
                try:
                    with open(final_report_txt_path, "w", encoding="utf-8") as f_txt:
                        f_txt.write(final_report_content_to_save)
                    logger.info(f"Итоговый отчет сохранен в TXT: {final_report_txt_path}")
                    st.success(f"✅ Итоговый отчет сохранен в TXT: {final_report_txt_path}")
                except Exception as e_txt:
                    error_msg_txt = f"❌ Ошибка при сохранении итогового отчета в TXT: {e_txt}"
                    logger.error(error_msg_txt)
                    st.error(error_msg_txt)

                # Сохранение в .docx
                try:
                    from docx import Document # Импортируем внутри блока try
                    doc = Document()
                    # Добавляем заголовок
                    doc.add_heading('Итоговый аналитический отчет', level=1)
                    # Добавляем содержимое отчета построчно
                    for line in final_report_content_to_save.splitlines():
                        if line.strip():
                            doc.add_paragraph(line)
                    doc.save(final_report_docx_path)
                    logger.info(f"Итоговый отчет сохранен в DOCX: {final_report_docx_path}")
                    st.success(f"✅ Итоговый отчет сохранен в DOCX: {final_report_docx_path}")
                except ImportError as e_import:
                    error_msg_import = f"❌ Библиотека python-docx не установлена. Итоговый отчет в DOCX не сохранен: {e_import}"
                    logger.error(error_msg_import)
                    st.error(error_msg_import)
                except Exception as e_docx:
                    error_msg_docx = f"❌ Ошибка при сохранении итогового отчета в DOCX: {e_docx}"
                    logger.error(error_msg_docx)
                    st.error(error_msg_docx)

                # Сообщение об успешном завершении анализа и сохранении результатов
                st.success("✅ Анализ завершён! Результаты сохранены.")

            except Exception as e: # Этот except ловит ошибки, возникающие при генерации отчета или его сохранении
                error_msg = f"❌ Ошибка при генерации отчёта: {e}"
                logger.error(error_msg, exc_info=True) # Добавляем exc_info для более подробного лога
                st.session_state["final_report"] = error_msg # Записываем ошибку в состояние
                st.error(error_msg) # Показываем ошибку пользователю

        # --- БЛОК СОХРАНЕНИЯ В report.txt УДАЛЕН, так как он дублирует сохранение выше ---
        # (Удалите или закомментируйте оригинальный блок сохранения report.txt здесь)
        # try:
        #     report_path = os.path.join(output_dir, "report.txt")
        #     report_content_to_save = st.session_state["final_report"] # <-- Эта строка вызывала ошибку
        #     if not isinstance(report_content_to_save, str):
        #          report_content_to_save = str(convert_numpy_types(report_content_to_save))
        #     with open(report_path, "w", encoding="utf-8") as f:
        #         f.write(report_content_to_save)
        #     st.session_state["report_path"] = report_path
        #     st.success("✅ Анализ завершён! Результаты сохранены.")
        # except Exception as e:
        #     st.error(f"❌ Ошибка при сохранении отчёта: {e}")

# -----------------------------------------
# Отображение результатов
# -----------------------------------------
st.header("4. Результаты анализа")
if "final_report" in st.session_state:
    with st.expander("📋 Структура данных (анализ LLM)"):
        st.text(st.session_state.get("data_structure", "Не определена"))
    with st.expander("📏 План метрик (JSON)"):
        st.code(st.session_state.get("metrics_plan", "{}"), language="json")
    with st.expander("💻 Сгенерированный код для расчёта"):
        st.code(st.session_state.get("calculation_code", "# Код не был сгенерирован."), language="python")
    st.write("#### 📈 Рассчитанные метрики:")
    st.code(st.session_state.get("metrics_results_raw", "Не рассчитано"), language="text")
    with st.expander("🔍 Анализ метрик"):
        analysis_summary_to_display = st.session_state.get("analysis_summary", "Анализ не выполнен.")
        if not isinstance(analysis_summary_to_display, str):
            analysis_summary_to_display = str(convert_numpy_types(analysis_summary_to_display))
        st.write(analysis_summary_to_display)
    with st.expander("🎨 Сгенерированный код визуализации"):
        st.code(st.session_state.get("viz_code", "# Код не сгенерирован"), language="python")
    st.write("#### 📄 Итоговый отчёт:")
    report_to_display = st.session_state["final_report"]
    if not isinstance(report_to_display, str):
        report_to_display = str(convert_numpy_types(report_to_display))
    st.text_area("", report_to_display, height=400)
    # # Кнопка скачивания отчёта
    # report_bytes = BytesIO(report_to_display.encode("utf-8"))
    # st.download_button(
    #     label="📥 Скачать отчёт (txt)",
    #     data=report_bytes,
    #     file_name="report.txt",
    #     mime="text/plain"
    # )
    # Отображение графиков
    st.write("#### 📊 Графики:")
    try:
        output_dir = st.session_state.get("output_dir", ".")
        plot_files = [f for f in os.listdir(output_dir) if f.startswith("plot_") and f.endswith(".png")]
        if plot_files:
            cols = st.columns(2)
            for i, plot in enumerate(plot_files):
                with cols[i % 2]:
                    st.image(os.path.join(output_dir, plot), caption=plot, use_container_width=True)
        else:
            st.info("Графики не найдены в указанной директории.")
    except Exception as e:
        st.error(f"Ошибка при отображении графиков: {e}")
