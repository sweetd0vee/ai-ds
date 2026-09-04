"""Экспорт графиков в DOCX с изображениями и выводами."""

from __future__ import annotations

import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from .analysis_export import ACCENT, BODY, MUTED, _add_body_paragraph, _set_run_font

logger = logging.getLogger(__name__)

MAX_IMAGE_WIDTH = Inches(6.2)
MAX_IMAGE_PX = 1280


def _image_for_docx(image_path: Path) -> str | BytesIO:
    """Уменьшить PNG перед вставкой — быстрее сборка DOCX и меньше файл."""
    try:
        from PIL import Image

        with Image.open(image_path) as img:
            if img.width <= MAX_IMAGE_PX:
                return str(image_path)
            ratio = MAX_IMAGE_PX / img.width
            resized = img.resize(
                (MAX_IMAGE_PX, max(1, int(img.height * ratio))),
                Image.Resampling.LANCZOS,
            )
            buf = BytesIO()
            resized.save(buf, format="PNG", optimize=True)
            buf.seek(0)
            return buf
    except Exception:
        logger.debug("Using original image for DOCX: %s", image_path, exc_info=True)
        return str(image_path)


def build_plots_docx(
    plot_details: list[dict],
    plots_dir: Path,
    output_path: Path,
    *,
    source_file: str = "",
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run("Отчёт по визуализациям")
    _set_run_font(title_run, size=20, bold=True, color=ACCENT)

    meta_parts = []
    if source_file:
        meta_parts.append(f"Файл: {source_file}")
    meta_parts.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    meta_parts.append(f"Графиков: {len(plot_details)}")
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_after = Pt(14)
    meta_run = meta_para.add_run(" · ".join(meta_parts))
    _set_run_font(meta_run, size=10, color=MUTED)

    if not plot_details:
        _add_body_paragraph(doc, "Графики не были построены для этой задачи.")
    else:
        for index, item in enumerate(plot_details, start=1):
            filename = item.get("filename", "")
            title = item.get("title") or filename or f"График {index}"
            conclusion = item.get("conclusion") or ""

            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(16 if index > 1 else 4)
            heading.paragraph_format.space_after = Pt(8)
            heading_run = heading.add_run(f"{index}. {title}")
            _set_run_font(heading_run, size=14, bold=True, color=BODY)

            image_path = plots_dir / filename
            if image_path.exists():
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.paragraph_format.space_after = Pt(10)
                run = pic_para.add_run()
                run.add_picture(_image_for_docx(image_path), width=MAX_IMAGE_WIDTH)
            else:
                missing = doc.add_paragraph()
                missing_run = missing.add_run(f"[Изображение не найдено: {filename}]")
                _set_run_font(missing_run, size=10, italic=True, color=MUTED)

            if filename:
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.space_after = Pt(10)
                cap_run = caption.add_run(filename)
                _set_run_font(cap_run, size=9, color=MUTED)

            label_para = doc.add_paragraph()
            label_para.paragraph_format.space_after = Pt(4)
            label_run = label_para.add_run("Выводы")
            _set_run_font(label_run, size=12, bold=True, color=ACCENT)

            _add_body_paragraph(doc, conclusion or "Интерпретация недоступна.")

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    footer_run = footer.add_run("Сформировано автоматически · Электронный Датасаентист")
    _set_run_font(footer_run, size=9, italic=True, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def ensure_plots_report_docx(
    output_dir: Path,
    plot_files: list[str],
    *,
    plot_details: list[dict] | None = None,
    source_file: str = "",
    correlations: dict | None = None,
    viz_output: str = "",
    dataset_path: str | None = None,
    datetime_candidates: list | None = None,
) -> Path:
    """Собрать DOCX один раз и переиспользовать готовый файл при скачивании."""
    output_path = output_dir / "plots_report.docx"
    if output_path.exists() and output_path.stat().st_size > 0:
        return output_path

    details = plot_details
    if not details:
        from .loaders import load_dataframe
        from .plot_insights import rebuild_plot_details
        from .preprocess import preprocess_dates_based_on_llm

        df = None
        if dataset_path:
            try:
                df = load_dataframe(dataset_path)
                df = preprocess_dates_based_on_llm(df, datetime_candidates or [])
            except Exception:
                logger.warning("Could not load dataset for plots_report.docx", exc_info=True)

        details = rebuild_plot_details(
            plot_files,
            df=df,
            correlations=correlations,
            viz_output=viz_output,
        )

    build_plots_docx(details, output_dir, output_path, source_file=source_file)
    return output_path
