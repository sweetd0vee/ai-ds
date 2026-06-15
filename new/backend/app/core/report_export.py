"""Экспорт итогового отчёта в DOCX с оформлением."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from .analysis_export import (
    ACCENT,
    BODY,
    MUTED,
    _add_body_paragraph,
    _add_markdown_runs,
    _add_subheading_paragraph,
    _normalize_text,
    _render_block,
    _set_run_font,
)

SECTION_RE = re.compile(r"^(\d+)\.\s+(.+)$")
BULLET_RE = re.compile(r"^•\s*(.+)$")
KV_RE = re.compile(r"^([^:]{2,48}):\s*(.+)$")


def _split_report_sections(text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_title: str | None = None
    body_lines: list[str] = []

    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        if stripped == "ИТОГОВЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ" or re.fullmatch(r"=+", stripped):
            continue
        if stripped.startswith("— Отчёт сформирован"):
            continue

        match = SECTION_RE.match(stripped)
        if match:
            if current_title is not None:
                sections.append((current_title, "\n".join(body_lines).strip()))
            current_title = match.group(2).strip()
            body_lines = []
        else:
            body_lines.append(line)

    if current_title is not None:
        sections.append((current_title, "\n".join(body_lines).strip()))
    return sections


def _add_section_heading(doc: Document, title: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(title)
    _set_run_font(run, size=14, bold=True, color=ACCENT)


def _add_kv_line(doc: Document, key: str, value: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    key_run = para.add_run(f"{key}: ")
    _set_run_font(key_run, bold=True, color=BODY)
    _add_markdown_runs(para, value)


def _add_bullet_line(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(3)
    _add_markdown_runs(para, text.strip())


def _clean_body(body: str) -> str:
    if "---JSON---" in body:
        body = body.split("---JSON---", 1)[0]
    return body.strip()


def _render_lines(doc: Document, body: str, *, bullets_only: bool = False):
    body = _clean_body(body)
    if not body:
        _add_body_paragraph(doc, "Данные недоступны.")
        return

    for line in body.split("\n"):
        stripped = line.strip()
        if not stripped or stripped.startswith("— "):
            continue

        bullet = BULLET_RE.match(stripped)
        if bullet:
            _add_bullet_line(doc, bullet.group(1))
            continue

        if bullets_only:
            _add_bullet_line(doc, stripped)
            continue

        kv = KV_RE.match(stripped)
        if kv and not stripped.startswith("http"):
            _add_kv_line(doc, kv.group(1).strip(), kv.group(2).strip())
            continue

        if stripped.isupper() and len(stripped) < 72 and "  " not in stripped:
            _add_subheading_paragraph(doc, stripped)
            continue

        _add_body_paragraph(doc, stripped)


def _render_section(doc: Document, title: str, body: str):
    upper = title.upper()

    if "ИНТЕРПРЕТАЦИЯ" in upper or "АНАЛИЗ" in upper:
        text = _normalize_text(_clean_body(body))
        blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]
        if not blocks:
            _add_body_paragraph(doc, "Анализ недоступен.")
        else:
            for block in blocks:
                _render_block(doc, block)
        return

    if any(k in upper for k in ("РЕКОМЕНДАЦИИ", "ВИЗУАЛИЗАЦ", "МЕТРИК")):
        _render_lines(doc, body, bullets_only="ВИЗУАЛИЗАЦ" in upper)
        return

    blocks = [b.strip() for b in re.split(r"\n\s*\n", _clean_body(body)) if b.strip()]
    if not blocks:
        _render_lines(doc, body)
        return

    for block in blocks:
        lines = [ln.strip() for ln in block.split("\n") if ln.strip()]
        if len(lines) > 1 and all(BULLET_RE.match(ln) or ln.startswith("•") for ln in lines):
            for ln in lines:
                content = BULLET_RE.match(ln).group(1) if BULLET_RE.match(ln) else ln.lstrip("•").strip()
                _add_bullet_line(doc, content)
        elif len(lines) == 1:
            ln = lines[0]
            if BULLET_RE.match(ln):
                _add_bullet_line(doc, BULLET_RE.match(ln).group(1))
            else:
                kv = KV_RE.match(ln)
                if kv:
                    _add_kv_line(doc, kv.group(1).strip(), kv.group(2).strip())
                elif ln.isupper() and len(ln) < 72:
                    _add_subheading_paragraph(doc, ln)
                else:
                    _add_body_paragraph(doc, ln)
        else:
            _render_lines(doc, block)


def build_report_docx(
    report_text: str,
    output_path: Path,
    *,
    source_file: str = "",
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run("Итоговый аналитический отчёт")
    _set_run_font(title_run, size=20, bold=True, color=ACCENT)

    meta_parts = []
    if source_file:
        meta_parts.append(f"Файл: {source_file}")
    meta_parts.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(14)
    meta_run = meta_para.add_run(" · ".join(meta_parts))
    _set_run_font(meta_run, size=10, color=MUTED)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(6)
    div_run = divider.add_run("─" * 48)
    _set_run_font(div_run, size=9, color=MUTED)

    sections = _split_report_sections(report_text or "")
    if not sections:
        _add_body_paragraph(doc, report_text or "Отчёт недоступен.")
    else:
        for title, body in sections:
            _add_section_heading(doc, title)
            _render_section(doc, title, body)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    footer_run = footer.add_run("Сформировано автоматически · Электронный Датасаентист")
    _set_run_font(footer_run, size=9, italic=True, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
