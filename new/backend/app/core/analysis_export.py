"""Экспорт текстового анализа метрик в DOCX с оформлением."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0x1D, 0x4E, 0xD8)
MUTED = RGBColor(0x6B, 0x72, 0x80)
BODY = RGBColor(0x1F, 0x29, 0x37)

FEATURE_RE = re.compile(r"^\*\*([^*]+)\*\*\s*[—–:-]\s*(.+)$", re.DOTALL)
SUBHEADING_RE = re.compile(r"^([А-ЯA-ZЁ][^:]{1,48}):\s*(.*)$", re.DOTALL)


def _set_run_font(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    try:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        r_fonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass


def _normalize_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").strip()
    text = re.sub(r"  \n", "\n", text)
    return text


def _add_markdown_runs(paragraph, text: str, *, base_size=11, base_color=BODY):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, size=base_size, bold=True, color=ACCENT)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, size=base_size, color=base_color)


def _configure_body_paragraph(para, *, space_after=8, bullet=False):
    if not bullet:
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.25
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_body_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    _configure_body_paragraph(para)
    _add_markdown_runs(para, text)
    return para


def _add_subheading_paragraph(doc: Document, title: str, rest: str = ""):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    title_run = para.add_run(title + (":" if not title.endswith(":") else ""))
    _set_run_font(title_run, size=12, bold=True, color=BODY)
    if rest:
        para.add_run(" ")
        _add_markdown_runs(para, rest)
    return para


def _add_feature_item(doc: Document, name: str, description: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.2
    name_run = para.add_run(name)
    _set_run_font(name_run, bold=True, color=ACCENT)
    sep = para.add_run(" — ")
    _set_run_font(sep, color=MUTED)
    _add_markdown_runs(para, description.rstrip("."))
    return para


def _split_feature_lines(block: str) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []
    for line in block.split("\n"):
        line = line.strip()
        if not line:
            continue
        match = FEATURE_RE.match(line)
        if match:
            items.append((match.group(1).strip(), match.group(2).strip()))
        else:
            chunks = re.split(r"(?=\*\*[^*]+\*\*)", line)
            for chunk in chunks:
                chunk = chunk.strip()
                if not chunk:
                    continue
                match = FEATURE_RE.match(chunk)
                if match:
                    items.append((match.group(1).strip(), match.group(2).strip()))
    return items


def _render_block(doc: Document, block: str):
    block = block.strip()
    if not block:
        return

    features = _split_feature_lines(block)
    if features and len(features) >= 2:
        for name, desc in features:
            _add_feature_item(doc, name, desc)
        return

    if features and len(features) == 1 and "\n" not in block:
        name, desc = features[0]
        _add_feature_item(doc, name, desc)
        return

    sub = SUBHEADING_RE.match(block)
    if sub and "**" not in sub.group(1):
        title, rest = sub.group(1).strip(), sub.group(2).strip()
        if rest:
            _add_subheading_paragraph(doc, title, rest)
        else:
            _add_subheading_paragraph(doc, title)
        return

    for line in block.split("\n"):
        line = line.strip()
        if not line:
            continue
        match = FEATURE_RE.match(line)
        if match:
            _add_feature_item(doc, match.group(1).strip(), match.group(2).strip())
        else:
            _add_body_paragraph(doc, line)


def build_analysis_docx(
    analysis_text: str,
    output_path: Path,
    *,
    title: str = "Анализ рассчитанных метрик",
    source_file: str = "",
) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run(title)
    _set_run_font(title_run, size=20, bold=True, color=ACCENT)

    meta_parts = []
    if source_file:
        meta_parts.append(f"Файл: {source_file}")
    meta_parts.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    meta_para = doc.add_paragraph()
    meta_para.paragraph_format.space_after = Pt(14)
    meta_run = meta_para.add_run(" · ".join(meta_parts))
    _set_run_font(meta_run, size=10, color=MUTED)

    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(12)
    div_run = divider.add_run("─" * 48)
    _set_run_font(div_run, size=9, color=MUTED)

    text = _normalize_text(analysis_text)
    blocks = [b.strip() for b in re.split(r"\n\s*\n", text) if b.strip()]

    if not blocks:
        _add_body_paragraph(doc, "Анализ недоступен.")
    else:
        for block in blocks:
            _render_block(doc, block)

    footer = doc.add_paragraph()
    footer.paragraph_format.space_before = Pt(18)
    footer_run = footer.add_run("Сформировано автоматически · Электронный Датасаентист")
    _set_run_font(footer_run, size=9, italic=True, color=MUTED)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
