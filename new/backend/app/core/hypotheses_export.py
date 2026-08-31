"""Экспорт гипотез в DOCX и XLSX."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from .analysis_export import ACCENT, BODY, MUTED
from .hypotheses import format_hypotheses_text
from .structure_export import HEADER_FILL, HEADER_FONT, THIN_BORDER

_PRIORITY_COLORS = {
    "high": RGBColor(0xB9, 0x1C, 0x1C),
    "medium": RGBColor(0xB4, 0x5A, 0x00),
    "low": RGBColor(0x05, 0x76, 0x69),
}

_PRIORITY_FILLS = {
    "high": ("FEE2E2", "B91C1C"),
    "medium": ("FEF3C7", "B45309"),
    "low": ("D1FAE5", "047857"),
}

_XLSX_HEADERS = [
    "№",
    "Название",
    "Формулировка",
    "Основание",
    "Столбцы",
    "Как проверить",
    "Приоритет",
    "Тип",
]

_XLSX_WIDTHS = {
    1: 8,
    2: 28,
    3: 42,
    4: 42,
    5: 22,
    6: 42,
    7: 14,
    8: 22,
}


def filter_hypotheses_by_ids(hypotheses: list[dict], ids: list[int] | None) -> list[dict]:
    items = [h for h in (hypotheses or []) if isinstance(h, dict)]
    if ids is None:
        return items
    wanted = {int(i) for i in ids}
    selected: list[dict] = []
    seen: set[int] = set()
    for item in items:
        try:
            hid = int(item.get("id") or 0)
        except (TypeError, ValueError):
            continue
        if hid in wanted and hid not in seen:
            selected.append(item)
            seen.add(hid)
    return selected


def build_hypotheses_docx(
    hypotheses: list[dict],
    output_path: Path,
    *,
    source_file: str = "",
    raw_fallback: str = "",
):
    doc = Document()
    title = doc.add_heading("Гипотезы о данных", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        + (f" · Файл: {source_file}" if source_file else "")
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    if hypotheses:
        for item in hypotheses:
            heading = doc.add_heading(f"{item.get('id', '?')}. {item.get('title', 'Гипотеза')}", level=2)
            for run in heading.runs:
                run.font.color.rgb = ACCENT

            priority = item.get("priority", "medium")
            p_line = doc.add_paragraph()
            p_run = p_line.add_run(f"Приоритет: {item.get('priority_label', priority)}")
            p_run.bold = True
            p_run.font.color.rgb = _PRIORITY_COLORS.get(priority, BODY)

            for label, key in (
                ("Формулировка", "statement"),
                ("Основание", "rationale"),
                ("Как проверить", "verification"),
            ):
                value = item.get(key, "")
                if not value:
                    continue
                para = doc.add_paragraph()
                label_run = para.add_run(f"{label}: ")
                label_run.bold = True
                label_run.font.color.rgb = BODY
                para.add_run(value)

            columns = item.get("columns") or []
            if columns:
                para = doc.add_paragraph()
                label_run = para.add_run("Столбцы: ")
                label_run.bold = True
                para.add_run(", ".join(columns))

            kind_label = item.get("kind_label") or item.get("kind")
            if kind_label:
                para = doc.add_paragraph()
                label_run = para.add_run("Тип: ")
                label_run.bold = True
                para.add_run(str(kind_label))

            doc.add_paragraph()
    else:
        for line in (raw_fallback or format_hypotheses_text([])).splitlines():
            if line.strip():
                doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def build_hypotheses_xlsx(
    hypotheses: list[dict],
    output_path: Path,
    *,
    source_file: str = "",
    raw_fallback: str = "",
):
    wb = Workbook()
    ws = wb.active
    ws.title = "Гипотезы"

    wrap = Alignment(wrap_text=True, vertical="top")
    wrap_center = Alignment(wrap_text=True, vertical="top", horizontal="center")

    row = 1
    ws.cell(row, 1, "Гипотезы о данных").font = Font(bold=True, size=14, color="111827")
    row += 1

    meta_parts = [f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}"]
    if source_file:
        meta_parts.append(f"Файл: {source_file}")
    if hypotheses:
        meta_parts.append(f"Гипотез: {len(hypotheses)}")
    ws.cell(row, 1, " · ".join(meta_parts)).font = Font(size=10, color="6B7280")
    row += 2

    if hypotheses:
        for col_idx, title in enumerate(_XLSX_HEADERS, start=1):
            cell = ws.cell(row, col_idx, title)
            cell.fill = HEADER_FILL
            cell.font = HEADER_FONT
            cell.border = THIN_BORDER
            cell.alignment = Alignment(
                horizontal="center" if col_idx in (1, 7) else "left",
                vertical="center",
                wrap_text=True,
            )
        header_row = row
        row += 1

        for item in hypotheses:
            values = [
                item.get("id", ""),
                item.get("title", ""),
                item.get("statement", ""),
                item.get("rationale", ""),
                ", ".join(item.get("columns") or []),
                item.get("verification", ""),
                item.get("priority_label") or item.get("priority", ""),
                item.get("kind_label") or item.get("kind") or "",
            ]
            for col_idx, value in enumerate(values, start=1):
                cell = ws.cell(row, col_idx, value if value not in (None, "") else "—")
                cell.border = THIN_BORDER
                cell.alignment = wrap_center if col_idx in (1, 7) else wrap
                cell.font = Font(size=10, color="111827")

            priority = str(item.get("priority") or "medium")
            bg, fg = _PRIORITY_FILLS.get(priority, ("F3F4F6", "4B5563"))
            prio_cell = ws.cell(row, 7)
            prio_cell.fill = PatternFill("solid", fgColor=bg)
            prio_cell.font = Font(bold=True, color=fg, size=10)
            row += 1

        ws.freeze_panes = f"A{header_row + 1}"
        ws.auto_filter.ref = f"A{header_row}:H{max(header_row, row - 1)}"
        ws.row_dimensions[header_row].height = 22
    else:
        fallback = (raw_fallback or format_hypotheses_text([])).strip()
        cell = ws.cell(row, 1, fallback or "Гипотезы не сформулированы.")
        cell.alignment = wrap
        cell.font = Font(size=10, color="4B5563")
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)

    for col, width in _XLSX_WIDTHS.items():
        ws.column_dimensions[get_column_letter(col)].width = width

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
